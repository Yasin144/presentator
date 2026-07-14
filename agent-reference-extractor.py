import json
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def extract(path: Path):
    suffix = path.suffix.lower()
    text = ""
    metadata = {"sizeBytes": path.stat().st_size, "extension": suffix}

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        metadata["pages"] = len(reader.pages)
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        metadata["paragraphs"] = len(document.paragraphs)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    elif suffix in {".txt", ".md", ".csv", ".json", ".log", ".srt"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported document type: {suffix}")

    return {
        "ok": True,
        "name": path.name,
        "kind": "document",
        "text": text[:120000],
        "truncated": len(text) > 120000,
        "metadata": metadata,
    }


if __name__ == "__main__":
    try:
        print(json.dumps(extract(Path(sys.argv[1]).resolve()), ensure_ascii=False))
    except Exception as error:
        print(json.dumps({"ok": False, "error": f"{type(error).__name__}: {error}"}))
        raise SystemExit(1)
